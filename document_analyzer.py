import PyPDF2
import google.generativeai as genai
import textwrap
from datetime import datetime
import sys
import os
import traceback
import time

# Try to import docx library for processing DOCX files
try:
    import docx
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    print("Warning: python-docx not installed. DOCX support will be limited.")

# API keys for Gemini
API_KEYS = [
    "AIzaSyABP0FhpPcNotV7TqlUw38Qm0YpAovfoIY",
    "AIzaSyBzB-FbuQimtmUEoaXUwYdGoxUwTXvMO3I"
]
current_key_index = 0

# Get next API key for rotation
def get_next_api_key():
    global current_key_index
    current_key_index = (current_key_index + 1) % len(API_KEYS)
    return API_KEYS[current_key_index]

def extract_text_from_pdf(pdf_path):
    """Extracts text from a PDF file using PyPDF2 with improved text handling."""
    text = ""
    try:
        with open(pdf_path, 'rb') as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            num_pages = len(pdf_reader.pages)
            
            print("\n[PDF] Processing {} pages...".format(num_pages))
            for i, page in enumerate(pdf_reader.pages, 1):
                text += f"\n─── PAGE {i} ───\n"
                page_text = page.extract_text()
                if not page_text:
                    page_text = "[Empty or unreadable page]"
                text += page_text + "\n"
                print("Page {} extracted".format(i), end='\r')
                
    except Exception as e:
        print("\n[ERROR] Error reading PDF: {}".format(e))
        print(traceback.format_exc())
        return None
    
    print("\n[SUCCESS] Text extraction complete")
    return text

def extract_text_from_docx(docx_path):
    """Extracts text from a DOCX file."""
    if not DOCX_SUPPORT:
        print("\n[WARNING] python-docx not installed. Using fallback method for DOCX.")
        try:
            with open(docx_path, 'rb') as f:
                import re
                raw_text = f.read().decode('utf-8', errors='ignore')
                text = re.sub(r'[^\x20-\x7E\n\r\t]', '', raw_text)
                print("\n[SUCCESS] Text extraction complete (fallback method)")
                return text
        except Exception as e:
            print("\n[ERROR] Error reading DOCX: {}".format(e))
            print(traceback.format_exc())
            return None
    
    try:
        doc = docx.Document(docx_path)
        text = ""
        print("\n[DOCX] Processing DOCX document...")
        
        # Extract headers
        for i, para in enumerate(doc.paragraphs):
            text += para.text + "\n"
            
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
            text += "\n"
            
        print("\n[SUCCESS] Text extraction complete")
        return text
    except Exception as e:
        print("\n[ERROR] Error reading DOCX: {}".format(e))
        print(traceback.format_exc())
        return None

def generate_detailed_summary(text):
    """Generates a comprehensive formatted summary"""
    prompt = """
    Create a DETAILED professional summary of this document with these sections:
    
    ─── DOCUMENT SUMMARY ───
    • Title: [Identify document title]
    • Date: [Find publication/release date]
    
    ─── KEY INFORMATION ───
    • Primary Authors/Entities: [List all]
    • Main Purpose: [2-3 sentences]
    • Scope/Coverage: [What areas does it address?]
    . Summary in Brief
    
    3. KEY OBLIGATIONS: Bullet points of material obligations for each party
    4. TERM: Duration of the agreement and renewal terms
    5. TERMINATION: Conditions under which agreement can be terminated
    6. GOVERNING LAW: Jurisdiction specified
    7. INDEMNIFICATION: Any indemnity clauses
    8. LIABILITY LIMITATIONS: Caps on liability or damages
    9. CONFIDENTIALITY: Existence and scope of confidentiality provisions
    
    ─── ADDITIONAL DETAILS ───
    • Special Features: [Charts/Tables/References]
    
    Use proper line breaks between sections and maintain professional formatting.
    """
    
    return analyze_text_with_gemini(text, prompt)

def analyze_text_with_gemini(text, prompt):
    """Enhanced analysis with better error handling and API key rotation."""
    if not text or len(text.strip()) < 50:
        print("\n[WARNING] The extracted text is too short or empty for analysis.")
        return "ERROR: The document appears to be empty or could not be properly read."
    
    max_retries = len(API_KEYS) * 2  # Try each key twice
    retries = 0
    last_error = None
    
    while retries < max_retries:
        try:
            # Get current API key
            api_key = API_KEYS[current_key_index]
            
            print(f"\nAttempt {retries+1}: Using API key ending in ...{api_key[-4:]}")
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel('gemini-2.5-flash')
            
            # Format the prompt without using system roles
            # Merge instructions and prompt into a single user prompt
            combined_prompt = f"""You are a professional document analyst. Provide thorough, well-structured responses with clear section headings.

{prompt}

Document text:
{text}"""
            
            response = model.generate_content(
                combined_prompt,
                generation_config={
                    "temperature": 0.3,  # More factual responses
                    "top_p": 0.95
                }
            )
            
            # Format the response
            return textwrap.fill(response.text, width=80, replace_whitespace=False)
            
        except Exception as e:
            last_error = e
            print(f"Error with key {api_key[-4:]} and model gemini-2.5-flash: {str(e)[:200]}")
            
            # Rotate API key on error
            get_next_api_key()
            
            # Exponential backoff
            wait_time = 750 * (1.5 ** retries)
            if wait_time > 5000:
                wait_time = 5000
                
            print(f"Rate limit hit. Rotating API key and waiting {wait_time}ms before retry {retries+1}")
            time.sleep(wait_time / 1000)  # Convert to seconds
            
            retries += 1
    
    # If all retries failed
    error_msg = f"ERROR: Could not complete analysis after {max_retries} attempts: {str(last_error)[:200]}"
    print(f"\n[ERROR] {error_msg}")
    return error_msg

def chat_with_pdf(extracted_text):
    """Enhanced chat interface with conversation history and API key rotation."""
    print("\n[CHAT] PDF CHAT MODE (type 'exit' to end)")
    print("Tip: Ask about specific sections, data, or request analyses")
    
    chat_history = []
    while True:
        try:
            user_input = input("\nQuestion: ").strip()
            
            if user_input.lower() in ['exit', 'quit']:
                print("\n[CHAT] Chat session ended")
                break
                
            if not user_input:
                print("Please enter a question")
                continue
                
            # Add context from chat history
            context = "\nPrevious questions:\n" + "\n".join(chat_history[-3:]) if chat_history else ""
            full_prompt = f"{context}\nNew Question: {user_input}\n\nAnswer in detail with page references if possible:"
            
            print("[ANALYSIS] Analyzing...", end='\r')
            response = analyze_text_with_gemini(extracted_text, full_prompt)
            
            if response:
                print("\n" + "─"*50)
                print("[RESPONSE]:\n{}".format(response))
                print("─"*50)
                chat_history.append(f"Q: {user_input}\nA: {response[:200]}...")
            else:
                print("[WARNING] No response generated")
                
        except KeyboardInterrupt:
            print("\nSession interrupted")
            break
        except Exception as e:
            print("\nError: {}".format(e))
            print(traceback.format_exc())

def process_text_file(file_path):
    """Process text file and return its content."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        print("\n[SUCCESS] Text extraction complete")
        return text
    except UnicodeDecodeError:
        # Try alternative encoding if UTF-8 fails
        try:
            with open(file_path, 'r', encoding='latin-1') as f:
                text = f.read()
            print("\n[SUCCESS] Text extraction complete (using latin-1 encoding)")
            return text
        except Exception as e:
            print("\n[ERROR] Error reading text file: {}".format(e))
            print(traceback.format_exc())
            return None
    except Exception as e:
        print("\n[ERROR] Error reading text file: {}".format(e))
        print(traceback.format_exc())
        return None

if __name__ == "__main__":
    # Check command-line arguments
    if len(sys.argv) < 2:
        print("[ERROR] Please provide a file path as an argument")
        print("Usage: python document_analyzer.py <file_path> [--no-interactive] [--capture-text]")
        sys.exit(1)
    
    file_path = sys.argv[1]
    
    # Parse flags
    no_interactive = '--no-interactive' in sys.argv
    capture_text = '--capture-text' in sys.argv
    
    if not os.path.exists(file_path):
        print("[ERROR] File not found at path: {}".format(file_path))
        sys.exit(1)
    
    print("\n" + "="*50)
    print("[FILE] Analyzing: {}".format(os.path.basename(file_path)))
    print("[TIME] Started: {}".format(datetime.now().strftime('%Y-%m-%d %H:%M:%S')))
    print("="*50)
    
    # Process the document based on file type
    file_ext = os.path.splitext(file_path)[1].lower()
    
    extracted_text = None
    
    try:
        if file_ext == '.pdf':
            extracted_text = extract_text_from_pdf(file_path)
        elif file_ext == '.txt':
            extracted_text = process_text_file(file_path)
        elif file_ext == '.docx':
            extracted_text = extract_text_from_docx(file_path)
        else:
            print("\n[ERROR] Unsupported file format: {}".format(file_ext))
            print("Currently supporting: PDF, TXT, DOCX")
            sys.exit(1)
    except Exception as e:
        print("\n[ERROR] Error processing file: {}".format(e))
        print(traceback.format_exc())
        sys.exit(1)
    
    if extracted_text:
        # Generate comprehensive summary
        print("\n[ANALYSIS] Generating professional summary...")
        summary = generate_detailed_summary(extracted_text)
        
        if summary:
            print("\n" + "="*50)
            print("[DOCUMENT SUMMARY]")
            print("="*50)
            print(summary)
            print("\n" + "="*50)
            
            # Output for API consumption
            print("\n─── ANALYSIS OUTPUT FOR API ───")
            print(summary)
            
            # Output document text for Q&A if requested
            if capture_text:
                print("\n─── DOCUMENT TEXT FOR QA ───")
                print(extracted_text)
            
            # Skip interactive mode when specified
            if not no_interactive:
                # Start interactive chat in command-line mode only
                chat_with_pdf(extracted_text)
        else:
            error_msg = "[ERROR] Failed to generate summary"
            print(error_msg)
            print("\n─── ANALYSIS OUTPUT FOR API ───")
            print(error_msg)
    else:
        error_msg = "[ERROR] Failed to process document"
        print(error_msg)
        print("\n─── ANALYSIS OUTPUT FOR API ───")
        print(error_msg)